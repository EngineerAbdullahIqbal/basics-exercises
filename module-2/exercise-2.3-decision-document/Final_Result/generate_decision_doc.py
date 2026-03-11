#!/usr/bin/env python3
"""Generate a formal Word decision document for career transition to AI/ML."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def create_decision_document():
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # =====================
    # COVER HEADER
    # =====================
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('CAREER TRANSITION DECISION DOCUMENT\n')
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Transitioning from Web Development to AI/ML Engineering\n')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(102, 102, 102)
    
    # Meta info
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run('Prepared: March 2026\n')
    meta.add_run('Decision Owner: Senior Web Developer\n')
    meta.add_run('Document Version: 1.0')
    
    doc.add_page_break()
    
    # =====================
    # SECTION 1: DECISION STATEMENT
    # =====================
    
    heading = doc.add_heading('1. Decision Statement', level=1)
    heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    
    p = doc.add_paragraph()
    p.add_run('Given $30,000 in savings, $2,500/month in fixed obligations, a supportive-but-risk-conscious partner, and 7 years of senior web development experience at $95K — which transition path into AI/ML engineering maximises long-term career and financial outcomes while keeping short-term financial risk within acceptable bounds?')
    p.style.font.bold = True
    p.style.font.size = Pt(12)
    
    # =====================
    # SECTION 2: CONTEXT & BACKGROUND
    # =====================
    
    heading = doc.add_heading('2. Context & Background', level=1)
    heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    
    doc.add_paragraph('Current Profile:', style='Heading 3')
    
    items = [
        ('Current Role:', 'Senior Web Developer (7 years experience)'),
        ('Current Salary:', '$95,000/year'),
        ('AI/ML Experience:', 'Self-taught through online courses and small personal projects; no professional AI experience'),
        ('Location:', 'Mid-sized city with limited but existing AI job market'),
        ('Financial Position:', '$30,000 in savings; $2,500/month fixed obligations'),
        ('Personal Situation:', 'Partner is supportive but risk-conscious regarding financial decisions'),
    ]
    
    for label, value in items:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(label).bold = True
        p.add_run(value)
    
    doc.add_paragraph('Motivation:', style='Heading 3')
    doc.add_paragraph(
        'AI/ML engineering represents significant career growth potential, higher long-term earning '
        'capacity, and the opportunity to work on cutting-edge technology. The field is experiencing '
        'unprecedented demand across all sectors, with companies integrating machine learning into '
        'their core products and services.'
    )
    
    # =====================
    # SECTION 3: ASSUMPTIONS & CONSTRAINTS
    # =====================
    
    heading = doc.add_heading('3. Assumptions & Constraints', level=1)
    heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    
    doc.add_paragraph('Hard Constraints:', style='Heading 3')
    
    constraints = [
        'Savings runway: $30,000 total savings provides 12 months of runway at current $2,500/month burn rate (with no income)',
        'Monthly obligations: $2,500/month fixed costs (rent, loans, living expenses) must be covered regardless of employment status',
        "Partner's risk tolerance: Supportive of career change but concerned about financial stability; prefers avoiding high-risk scenarios that could jeopardise household security",
        'Geographic limits: Mid-sized city with limited AI job market; may require remote work or relocation for optimal opportunities',
        'Current market conditions: AI hiring remains strong but increasingly competitive for entry-level roles; many positions prefer advanced degrees or demonstrated professional experience',
        'Time availability: Can dedicate 10-15 hours/week to learning if maintaining current employment',
    ]
    
    for constraint in constraints:
        doc.add_paragraph(constraint, style='List Bullet')
    
    doc.add_paragraph('Key Assumptions:', style='Heading 3')
    
    assumptions = [
        'Web development skills remain marketable throughout transition period',
        'AI/ML job market will remain viable over the next 12-24 months',
        'Willingness to potentially accept initial salary reduction to enter AI/ML field',
        'Ability to learn technical material independently with structured resources',
    ]
    
    for assumption in assumptions:
        doc.add_paragraph(assumption, style='List Bullet')
    
    # =====================
    # SECTION 4: THREE OPTIONS ANALYSED
    # =====================
    
    heading = doc.add_heading('4. Three Options Analysed', level=1)
    heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    
    # ---- Option 1 ----
    doc.add_heading('Option 1: Full Career Pause — 6-Month Intensive Bootcamp', level=2)
    
    doc.add_paragraph('Description:', style='Heading 3')
    doc.add_paragraph(
        'Quit current job and enroll in an intensive 6-month AI/ML bootcamp or structured program. '
        'Commit to full-time study with the goal of rapid transition into entry-level AI/ML roles '
        'upon completion.'
    )
    
    doc.add_paragraph('Timeline:', style='Heading 3')
    doc.add_paragraph('6 months full-time study + 1-3 months job search = 7-9 months total transition', style='List Bullet')
    
    doc.add_paragraph('Total Financial Cost:', style='Heading 3')
    costs = [
        'Bootcamp tuition: ~$15,000',
        'Lost income (6 months at $95K): ~$47,500',
        'Living expenses during study (6 × $2,500): $15,000 (covered by savings)',
        'Total direct + opportunity cost: ~$62,500',
        'Net impact on savings: -$30,000 (depletes entire savings)',
    ]
    for cost in costs:
        doc.add_paragraph(cost, style='List Bullet')
    
    doc.add_paragraph('Learning Depth Achieved:', style='Heading 3')
    doc.add_paragraph(
        'HIGH — Full immersion provides comprehensive coverage of ML fundamentals, deep learning, '
        'and practical projects. Structured curriculum ensures no gaps. Access to instructors and '
        'mentors accelerates learning. Portfolio development is built into program.',
        style='List Bullet'
    )
    
    doc.add_paragraph('Job Market Positioning on Completion:', style='Heading 3')
    doc.add_paragraph(
        'Entry-level AI/ML Engineer or Junior Data Scientist roles. Bootcamp credential provides '
        'signaling value but competes with candidates holding advanced degrees. Strong portfolio '
        'critical for differentiation. Career services support job search but placement not guaranteed.',
        style='List Bullet'
    )
    
    # ---- Option 2 ----
    doc.add_heading('Option 2: Gradual Transition — Keep Job, Study 10-15 Hours/Week', level=2)
    
    doc.add_paragraph('Description:', style='Heading 3')
    doc.add_paragraph(
        'Maintain current web development role while systematically building AI/ML skills through '
        'online courses, personal projects, and self-directed study. Dedicate evenings and weekends '
        'to learning. Target AI/ML roles after building sufficient competency and portfolio.',
        style='List Bullet'
    )
    
    doc.add_paragraph('Timeline:', style='Heading 3')
    doc.add_paragraph('12-18 months part-time study while employed + 1-3 months job search = 13-21 months total transition', style='List Bullet')
    
    doc.add_paragraph('Total Financial Cost:', style='Heading 3')
    costs = [
        'Online courses and resources: $1,000-$3,000',
        'Lost income: $0 (maintain current salary)',
        'Living expenses: Covered by ongoing salary',
        'Total direct cost: ~$2,000',
        'Opportunity cost: Delayed higher earnings for 12-18 months',
    ]
    for cost in costs:
        doc.add_paragraph(cost, style='List Bullet')
    
    doc.add_paragraph('Learning Depth Achieved:', style='Heading 3')
    doc.add_paragraph(
        'MEDIUM to HIGH — Self-directed learning allows customization but requires discipline. '
        'Risk of knowledge gaps without structured curriculum. Can achieve strong practical skills '
        'through projects. May lack theoretical depth compared to intensive programs.',
        style='List Bullet'
    )
    
    doc.add_paragraph('Job Market Positioning on Completion:', style='Heading 3')
    doc.add_paragraph(
        'Entry-level AI/ML Engineer, ML Engineer, or AI-adjacent roles. No employment gap on resume. '
        'Portfolio of personal projects demonstrates commitment. May need to overcome perception of '
        'less rigorous training compared to degree holders. Strong web development background is asset.',
        style='List Bullet'
    )
    
    # ---- Option 3 ----
    doc.add_heading('Option 3: AI-Adjacent Bridge Role — Leverage Web Skills While Learning AI', level=2)
    
    doc.add_paragraph('Description:', style='Heading 3')
    doc.add_paragraph(
        'Transition to a role that bridges web development and AI/ML, such as building AI-powered '
        'products, ML infrastructure, or MLOps. Leverage existing web development expertise to enter '
        'AI-adjacent positions, then deepen ML knowledge on the job. Target roles include ML Platform '
        'Engineer, AI Product Developer, or Full-Stack Engineer on AI teams.',
        style='List Bullet'
    )
    
    doc.add_paragraph('Timeline:', style='Heading 3')
    doc.add_paragraph('6-9 months upskilling + job search = 6-12 months to first AI-adjacent role; 18-24 months to core ML role', style='List Bullet')
    
    doc.add_paragraph('Total Financial Cost:', style='Heading 3')
    costs = [
        'Learning resources (MLOps, cloud ML, AI integration): $500-$2,000',
        'Lost income: $0 (may maintain or slightly reduce salary during transition)',
        'Living expenses: Covered by ongoing salary',
        'Total direct cost: ~$1,500',
        'Potential salary adjustment: May take -10% to +0% change initially; long-term upside significant',
    ]
    for cost in costs:
        doc.add_paragraph(cost, style='List Bullet')
    
    doc.add_paragraph('Learning Depth Achieved:', style='Heading 3')
    doc.add_paragraph(
        'MEDIUM initially, HIGH over time — Learn practical AI/ML through production work. May lack '
        'theoretical depth initially but gain real-world deployment experience. Can supplement with '
        'courses to build theory. On-the-job learning is highly relevant and immediately applicable.',
        style='List Bullet'
    )
    
    doc.add_paragraph('Job Market Positioning on Completion:', style='Heading 3')
    doc.add_paragraph(
        'ML Engineer (infrastructure/deployment focus), AI Product Engineer, or MLOps Engineer. '
        'High demand for engineers who can productionize ML systems. Combines valuable web skills '
        'with emerging AI expertise. Clear pathway to core ML roles after 1-2 years experience.',
        style='List Bullet'
    )
    
    # =====================
    # SECTION 5: RISK ASSESSMENT TABLE
    # =====================
    
    heading = doc.add_heading('5. Risk Assessment Table', level=1)
    heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    
    # Create table
    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header row
    header_cells = table.rows[0].cells
    headers = ['Risk Factor', 'Option 1: Full Pause', 'Option 2: Gradual', 'Option 3: Bridge']
    for i, header_text in enumerate(headers):
        cell = header_cells[i]
        cell.text = header_text
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Shade header
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'D6EAF8')
        cell._tc.append(shading)
    
    # Data rows
    risk_data = [
        ('Financial Risk', 'HIGH', 'LOW', 'LOW'),
        ('Career Risk', 'HIGH', 'MEDIUM', 'LOW'),
        ('Learning Depth', 'HIGH', 'MEDIUM', 'MEDIUM'),
        ('Speed to Transition', 'HIGH (Fast)', 'LOW (Slow)', 'MEDIUM'),
        ('Relationship/Lifestyle Impact', 'HIGH (stress on partner)', 'MEDIUM (time trade-offs)', 'LOW'),
    ]
    
    # Color mapping
    color_map = {
        'HIGH': 'FADBD8',      # Light red
        'MEDIUM': 'FDEBD0',    # Light orange
        'LOW': 'D5F5E3',       # Light green
    }
    
    for row_idx, (risk_factor, opt1, opt2, opt3) in enumerate(risk_data, start=1):
        row = table.rows[row_idx]
        cells = row.cells
        
        cells[0].text = risk_factor
        cells[0].paragraphs[0].runs[0].bold = True
        
        for col_idx, value in enumerate([opt1, opt2, opt3], start=1):
            cells[col_idx].text = value
            cells[col_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Apply color shading
            color = color_map.get(value.split()[0], 'FFFFFF')
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), color)
            cells[col_idx]._tc.append(shading)
    
    # Add table note
    doc.add_paragraph('\nNote: "Speed to Transition" rated HIGH for faster, LOW for slower.')
    
    # =====================
    # SECTION 6: RECOMMENDED PATH FORWARD
    # =====================
    
    heading = doc.add_heading('6. Recommended Path Forward', level=1)
    heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    
    recommendation = doc.add_paragraph()
    recommendation.add_run('RECOMMENDATION: Option 2 (Gradual Transition) with Option 3 (Bridge Role) Elements').bold = True
    recommendation.add_run('\n\n').font.size = Pt(6)
    
    rationale = doc.add_paragraph()
    rationale.add_run('Rationale:').bold = True
    rationale_text = (
        'Given your constraints—particularly the $30,000 savings with $2,500/month obligations '
        'and a risk-conscious partner—Option 1 (Full Pause) presents unacceptable financial risk. '
        'Depleting your entire savings while losing $47,500+ in income would create significant '
        'household stress and leave no buffer for emergencies or extended job search.\n\n'
    )
    rationale.add_run(rationale_text)
    
    rationale_text2 = (
        'Option 2 (Gradual Transition) is the recommended primary path because it:\n'
    )
    rationale.add_run(rationale_text2).bold = True
    
    reasons = [
        'Maintains your $95K salary throughout the transition, eliminating financial risk',
        'Respects your partner\'s risk tolerance by avoiding drastic income disruption',
        'Provides 12-18 months to build genuine competency without pressure',
        'Allows you to test your interest in AI/ML before fully committing',
        'Keeps all options open—you can accelerate, pivot, or pause based on market conditions',
    ]
    
    for reason in reasons:
        doc.add_paragraph(reason, style='List Bullet')
    
    doc.add_paragraph('\nIntegration of Option 3 (Bridge Role):')
    doc.add_paragraph(
        'While pursuing Option 2, actively seek AI-adjacent opportunities at your current company '
        'or in your next role. This hybrid approach lets you:\n'
    )
    doc.add_paragraph('Build AI experience on the job while maintaining income', style='List Bullet')
    doc.add_paragraph('Position yourself for internal transfer to AI teams', style='List Bullet')
    doc.add_paragraph('Accelerate transition timeline if bridge opportunities arise', style='List Bullet')
    
    doc.add_paragraph('\nAddressing Partner\'s Financial Concerns:')
    doc.add_paragraph(
        'This recommendation directly addresses your partner\'s concerns by maintaining household '
        'financial stability throughout the transition. You keep your income, preserve your $30,000 '
        'savings as an emergency buffer, and invest only $2,000-$3,000 in learning resources—a '
        'fraction of the $62,500+ cost of Option 1. If the AI market shifts or the transition takes '
        'longer than expected, you have not jeopardised your financial foundation.'
    )
    
    # =====================
    # SECTION 7: NEXT STEPS
    # =====================
    
    heading = doc.add_heading('7. Next Steps', level=1)
    heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    
    doc.add_paragraph('Five Concrete Actions:')
    
    next_steps = [
        ('Enroll in foundational ML course (Andrew Ng\'s Machine Learning on Coursera or equivalent)', 'You', 'End of this week'),
        ('Research AI-adjacent roles and internal opportunities at current company', 'You', 'End of this week'),
        ('Schedule meeting with manager to discuss AI initiatives and potential involvement', 'You', 'Within 2 weeks'),
        ('Build first AI-powered project integrating ML into a web application (e.g., recommendation feature, chatbot)', 'You', 'Within 8 weeks'),
        ('Join AI/ML community and schedule 2 informational interviews with AI engineers', 'You', 'Within 4 weeks'),
    ]
    
    # Create table for next steps
    steps_table = doc.add_table(rows=6, cols=3)
    steps_table.style = 'Table Grid'
    
    # Header
    header_cells = steps_table.rows[0].cells
    headers = ['Action', 'Owner', 'Deadline']
    for i, header_text in enumerate(headers):
        cell = header_cells[i]
        cell.text = header_text
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'D6EAF8')
        cell._tc.append(shading)
    
    # Data
    for row_idx, (action, owner, deadline) in enumerate(next_steps, start=1):
        row = steps_table.rows[row_idx]
        row.cells[0].text = action
        row.cells[1].text = owner
        row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row.cells[2].text = deadline
        row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # =====================
    # FINAL DECISION SUMMARY
    # =====================
    
    doc.add_page_break()
    
    summary = doc.add_paragraph()
    summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    summary_text = (
        'DECISION SUMMARY: Pursue a gradual transition into AI/ML while maintaining current employment, '
        'investing 10-15 hours/week over 12-18 months in structured learning and portfolio development, '
        'while actively seeking AI-adjacent opportunities to accelerate the transition without '
        'compromising household financial security.'
    )
    run = summary.add_run(summary_text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    # Save document
    output_path = '/home/abdullahiqbal/Abdullah/Digita-FTEs/basics-exercises/module-2/exercise-2.3-decision-document/Career-Transition-Decision-Document.docx'
    doc.save(output_path)
    print(f'Document saved to: {output_path}')

if __name__ == '__main__':
    create_decision_document()
